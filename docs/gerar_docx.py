from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2.5)

# =========================================================================
# HELPERS
# =========================================================================
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x16, 0x21, 0x3e)
    return h

def add_table_row(table, cells_text, bold=False, header=False):
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        if bold or header:
            run.bold = True
        if header:
            shading = cell._element.get_or_add_tcPr()
            shading_elm = shading.makeelement(qn('w:shd'), {
                qn('w:fill'): '16213e',
                qn('w:val'): 'clear'
            })
            shading.append(shading_elm)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

def create_table(headers, rows):
    table = doc.add_table(rows=0, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(table, headers, header=True)
    for row_data in rows:
        add_table_row(table, row_data)
    return table

def add_para(text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

def add_code_block(lines):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)

# =========================================================================
# CAPA
# =========================================================================
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Relatório de Testes Unitários e\nCobertura de Código')
run.font.size = Pt(22)
run.bold = True
run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
run.font.name = 'Times New Roman'

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Projeto AutoJuri — Frontend de Automação de Contestações Jurídicas')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
run.font.name = 'Times New Roman'

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Disciplina: Teste e Qualidade de Software')
run.font.size = Pt(14)
run.bold = True
run.font.name = 'Times New Roman'

for _ in range(4):
    doc.add_paragraph()

info_lines = [
    'Aluno: Guilherme Maciel',
    '',
    'Repositório: API-DE-AUTOMA-AO-DE-CONTESTACAO',
    '',
    'Tecnologias: React 19 • Vite 7 • Vitest 3 • V8 Coverage',
    '',
    'Data: 12 de Abril de 2026',
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(12)
    run.font.name = 'Times New Roman'

# =========================================================================
# SEÇÃO 1 — PASTA/DIRETÓRIO CONTENDO OS TESTES
# =========================================================================
doc.add_page_break()

add_heading_styled('1. Pasta e Diretório Contendo os Testes', level=1)

add_heading_styled('1.1 Estrutura do Projeto', level=2)

add_para(
    'O projeto AutoJuri é uma aplicação web de automação de contestações jurídicas, '
    'construída com React 19 e Vite 7. A lógica de negócio testável está concentrada '
    'no diretório src/utils/, que contém funções puras responsáveis por validações, '
    'formatações e operações de persistência local.'
)

add_para(
    'Os testes unitários seguem a convenção de co-location, ficando no mesmo '
    'diretório dos módulos testados, com o sufixo .test.js:'
)

add_code_block([
    'Front comp/vite-project/',
    '  src/',
    '    utils/',
    '      validators.js          → Validação de e-mail, senha, CNJ e campos de auth',
    '      validators.test.js     → 65 testes',
    '      cases.js               → Geração de ID incremental de caso jurídico',
    '      cases.test.js          → 7 testes',
    '      html.js                → Escape de caracteres HTML (proteção XSS)',
    '      html.test.js           → 11 testes',
    '      files.js               → Validação de upload, normalização de nome',
    '      files.test.js          → 23 testes',
    '      storage.js             → Persistência de rascunho e sessão no localStorage',
    '      storage.test.js        → 15 testes',
])

add_heading_styled('1.2 Módulo Selecionado e Justificativa', level=2)

add_para(
    'O módulo src/utils/ foi selecionado por conter toda a lógica de negócio do '
    'frontend, incluindo:'
)

add_bullet('validators.js: Regras de validação de formulários de autenticação (e-mail, senha com requisitos de força, número de processo CNJ) e extração de mensagens de erro da API;')
add_bullet('files.js: Validação de upload de peças processuais (extensão, MIME type, limite de tamanho) e normalização de nomes para exportação;')
add_bullet('storage.js: Persistência segura de rascunhos e sessões no navegador, garantindo que dados sensíveis como tokens não sejam armazenados;')
add_bullet('cases.js: Geração de identificadores únicos para casos no dashboard;')
add_bullet('html.js: Sanitização de strings para prevenção de ataques XSS.')

add_para(
    'Essas funções são ideais para testes unitários por serem predominantemente '
    'funções puras (sem efeitos colaterais além do localStorage) e por representarem '
    'regras críticas de segurança e integridade de dados da aplicação.'
)

add_heading_styled('1.3 Ferramentas Utilizadas', level=2)

create_table(
    ['Ferramenta', 'Versão', 'Finalidade'],
    [
        ['Vitest', '3.2.4', 'Framework de testes unitários (compatível com Vite)'],
        ['@vitest/coverage-v8', '3.2.4', 'Provedor de cobertura de código (engine V8)'],
        ['Node.js', '24.x', 'Runtime de execução dos testes'],
    ]
)

add_heading_styled('1.4 Como Executar', level=2)

add_code_block([
    '# Instalar dependências',
    'npm install',
    '',
    '# Executar todos os testes',
    'npm test',
    '',
    '# Executar com relatório de cobertura',
    'npm run test:coverage',
    '',
    '# Modo watch (desenvolvimento)',
    'npm run test:watch',
])

# =========================================================================
# SEÇÃO 2 — RELATÓRIO DE COBERTURA DE CÓDIGO
# =========================================================================
doc.add_page_break()

add_heading_styled('2. Relatório de Cobertura de Código', level=1)

add_heading_styled('2.1 Resumo Geral', level=2)

add_para(
    'A cobertura foi gerada com o provedor V8 do Vitest, que instrumenta o código '
    'diretamente na engine JavaScript. Os resultados refletem a execução de 121 testes '
    'distribuídos em 5 arquivos de teste, todos passando com sucesso.'
)

create_table(
    ['Métrica', 'Resultado'],
    [
        ['Total de testes', '121'],
        ['Testes passando', '121 (100%)'],
        ['Testes falhando', '0'],
        ['Arquivos de teste', '5'],
        ['Tempo de execução', '1.34s'],
    ]
)

add_heading_styled('2.2 Cobertura por Arquivo', level=2)

create_table(
    ['Arquivo', 'Statements', 'Branches', 'Functions', 'Lines', 'Linhas Não Cobertas'],
    [
        ['validators.js', '100%', '100%', '100%', '100%', '—'],
        ['cases.js', '100%', '100%', '100%', '100%', '—'],
        ['html.js', '100%', '100%', '100%', '100%', '—'],
        ['files.js', '83,07%', '86,66%', '100%', '83,07%', '63–73'],
        ['storage.js', '97,14%', '80%', '100%', '97,14%', '16–17'],
        ['TOTAL', '94,69%', '91,76%', '100%', '94,69%', '—'],
    ]
)

add_heading_styled('2.3 Distribuição de Testes por Módulo', level=2)

create_table(
    ['Arquivo de Teste', 'Qtd. Testes', 'Cenários Cobertos'],
    [
        ['validators.test.js', '65', 'E-mail, senha, CNJ, campos de auth, erros de API'],
        ['files.test.js', '23', 'Extensão, MIME, tamanho, normalização, constantes'],
        ['storage.test.js', '15', 'Rascunho, sessão, JSON inválido, segurança'],
        ['html.test.js', '11', 'Escape de caracteres, XSS com script/onerror'],
        ['cases.test.js', '7', 'ID incremental, padding, ano fictício, formato inválido'],
    ]
)

add_heading_styled('2.4 Análise das Linhas Não Cobertas', level=2)

add_para('files.js — Linhas 63 a 73 (função readFileAsBase64)', bold=True)
add_para(
    'A função readFileAsBase64 utiliza a API FileReader do navegador para converter '
    'arquivos em base64. Como os testes rodam em ambiente Node.js (sem DOM real), não '
    'é possível instanciar um FileReader nativo. Os caminhos de null/undefined foram '
    'cobertos, mas o fluxo principal exigiria um ambiente jsdom ou happy-dom com mock '
    'de FileReader.'
)

add_para('storage.js — Linhas 16 a 17 (guarda typeof window)', bold=True)
add_para(
    'O branch typeof window === "undefined" protege o código contra execução em '
    'ambientes server-side (SSR). Como o ambiente de teste possui globalThis.window '
    'definido, o caminho do retorno antecipado não é atingido. Trata-se de uma guarda '
    'defensiva cujo teste exigiria manipulação do objeto global entre execuções.'
)

# =========================================================================
# SEÇÃO 3 — ANÁLISE CRÍTICA
# =========================================================================
doc.add_page_break()

add_heading_styled('3. Análise Crítica', level=1)

add_heading_styled('3.1 Qualidade dos Testes Implementados', level=2)

add_para(
    'Os 121 testes unitários foram estruturados seguindo o padrão AAA '
    '(Arrange–Act–Assert), com nomenclatura descritiva em português que facilita a '
    'compreensão do comportamento esperado de cada função. A suíte cobre três '
    'categorias fundamentais de cenários:'
)

add_bullet(
    'Cenários principais (happy path): Validação de e-mails válidos, senhas fortes, '
    'números CNJ no formato correto, upload de arquivos PDF/DOC/DOCX dentro do limite, '
    'persistência e recuperação de rascunhos e sessões;'
)
add_bullet(
    'Casos de erro: E-mails mal formatados, senhas que não atendem aos requisitos de '
    'força, JSON inválido retornado pela API, campos obrigatórios vazios, arquivos com '
    'extensão proibida ou MIME type incorreto;'
)
add_bullet(
    'Casos de limite (edge cases): Valores null e undefined em todas as funções, '
    'strings contendo apenas espaços, arquivo exatamente no limite de 10MB, senha no '
    'limite máximo de 128 caracteres, IDs com formato corrompido (NaN), proteção '
    'contra XSS com payloads reais (<script>, onerror).'
)

add_para(
    'Um ponto de destaque é o teste de segurança em persistSession, que verifica que '
    'tokens sensíveis não são persistidos no localStorage — um requisito crítico para '
    'uma aplicação jurídica que lida com dados de clientes.'
)

add_heading_styled('3.2 Avaliação da Cobertura Atingida', level=2)

add_para(
    'A cobertura geral de 94,69% de statements e 100% de functions indica que '
    'praticamente toda a lógica de negócio testável foi exercitada. Três dos cinco '
    'módulos atingiram 100% em todas as métricas (validators, cases, html), o que '
    'demonstra cobertura completa das regras mais críticas do sistema.'
)

add_para(
    'A cobertura de branches em 91,76% reflete a dificuldade técnica de simular '
    'certos ambientes no Node.js: a guarda typeof window em storage.js e o fluxo '
    'do FileReader em files.js. Essas lacunas são justificáveis e não representam '
    'risco real, pois os branches não cobertos são guardas defensivas de ambiente, '
    'não lógica de negócio.'
)

add_heading_styled('3.3 Lacunas Identificadas', level=2)

create_table(
    ['Lacuna', 'Impacto', 'Justificativa'],
    [
        ['readFileAsBase64 — fluxo principal não coberto', 'Médio',
         'Depende de FileReader (API de browser). Exigiria ambiente jsdom com mock, ou testes E2E.'],
        ['Guarda SSR (typeof window) não exercitada', 'Baixo',
         'Código defensivo para SSR. A aplicação é SPA e não usa renderização server-side.'],
        ['Componentes React não testados', 'Médio',
         'Fora do escopo (lógica de apresentação). Recomendado para fase futura com React Testing Library.'],
        ['Integração com Supabase não testada', 'Alto',
         'Comunicação com backend sem testes. Requer mocks ou ambiente de staging.'],
    ]
)

add_heading_styled('3.4 Riscos Identificados', level=2)

add_bullet(
    'Validação apenas no frontend: As validações de e-mail, senha e número CNJ '
    'existem no cliente, mas um atacante pode enviá-las diretamente para a API. Os '
    'testes confirmam que as regras do frontend estão corretas, mas a segurança real '
    'depende de validação duplicada no backend;'
)
add_bullet(
    'Arquivo renomeado com MIME vazio: O sistema aceita arquivos com MIME type vazio '
    'como fallback para navegadores que não preenchem o Content-Type. Isso é '
    'intencional (documentado no código), mas abre uma janela para envio de arquivos '
    'maliciosos. O teste documenta esse comportamento deliberado;'
)
add_bullet(
    'localStorage não criptografado: Rascunhos de contestações jurídicas são salvos '
    'em texto plano no localStorage. Embora a persistSession filtre tokens, o conteúdo '
    'do rascunho pode conter dados sigilosos do cliente.'
)

add_heading_styled('3.5 Possíveis Melhorias', level=2)

improvements = [
    'Adicionar ambiente jsdom/happy-dom ao Vitest para cobrir readFileAsBase64 e simular interações com DOM, atingindo cobertura próxima de 100%;',
    'Implementar testes de componentes React com @testing-library/react para cobrir os formulários de AuthModal e MainPanelSection, que consomem diretamente as funções validadoras;',
    'Adicionar testes de integração com mock do Supabase para validar o fluxo completo de autenticação e envio de contestações;',
    'Configurar threshold de cobertura no Vitest (ex: branches: 90) para impedir que futuras alterações reduzam a cobertura abaixo do mínimo aceitável;',
    'Incluir testes de mutação (ex: Stryker) para verificar se os testes realmente detectam bugs introduzidos artificialmente, além de simplesmente exercitar as linhas.',
]
for i, text in enumerate(improvements, 1):
    p = doc.add_paragraph()
    run = p.add_run(f'{i}. {text}')
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

add_heading_styled('3.6 Conclusão', level=2)

add_para(
    'A implementação de 121 testes unitários atingiu uma cobertura de 94,69% nos '
    'módulos de lógica de negócio, com 100% de cobertura de funções e cobertura '
    'completa nos três módulos mais críticos (validators, cases, html). Os testes '
    'cobrem cenários principais, casos de erro e edge cases relevantes, incluindo '
    'proteção contra XSS e validação de segurança na persistência de sessão.'
)

add_para(
    'As lacunas identificadas são justificáveis por limitações do ambiente de teste '
    '(ausência de DOM real) e não comprometem a confiabilidade da suíte. O projeto '
    'está em uma base sólida para evoluir com testes de componente e integração nas '
    'próximas iterações.'
)

# =========================================================================
# SALVAR
# =========================================================================
output_path = os.path.join(os.path.dirname(__file__), 'Relatorio_Testes_Qualidade_AutoJuri.docx')
doc.save(output_path)
print(f'Arquivo salvo em: {output_path}')
